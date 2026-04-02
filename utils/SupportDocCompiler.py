import requests
import io
from flask import send_file
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
from pdf2image import convert_from_bytes


# ── File type constants ───────────────────────────────────────────────────────

IMAGE_TYPES = {"image/jpeg", "image/jpg", "image/png", "image/gif", "image/webp"}
PDF_TYPE    = "application/pdf"
WORD_TYPES  = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# ── Colour palette ────────────────────────────────────────────────────────────

COLOR_PRIMARY   = RGBColor(0x1A, 0x3A, 0x5C)  # deep navy
COLOR_SECONDARY = RGBColor(0x2E, 0x74, 0xB5)  # corporate blue
COLOR_MID_GRAY  = RGBColor(0x59, 0x59, 0x59)
COLOR_DARK_GRAY = RGBColor(0x26, 0x26, 0x26)
COLOR_RED       = RGBColor(0xC0, 0x00, 0x00)

HEX_BLUE  = "2E74B5"
HEX_NAVY  = "1A3A5C"
HEX_GRAY  = "AAAAAA"
HEX_GOLD  = "BF9300"


# ── Data collectors ───────────────────────────────────────────────────────────

def collect_by_ipcr(ipcr_id):
    from models.PCR import Supporting_Document, IPCR
    try:
        ipcr = IPCR.query.get(ipcr_id)
        if not ipcr:
            return []
        return [d.to_dict() for d in ipcr.supporting_documents if d.status and d.isApproved == "valid"]
    except Exception as e:
        print("collect_by_ipcr error:", e)
        return []


def collect_by_department(dept_id):
    from models.PCR import Supporting_Document
    from models.System_Settings import System_Settings
    try:
        settings = System_Settings.get_default_settings()
        all_docs = Supporting_Document.query.filter_by(
            period=settings.current_period_id
        ).all()
        return [
            d.to_dict() for d in all_docs
            if str(d.ipcr.user.department.id) == str(dept_id) and d.status and d.isApproved == "valid"
        ]
    except Exception as e:
        print("collect_by_department error:", e)
        return []


# ── XML / layout helpers ──────────────────────────────────────────────────────

def _set_default_font(doc, font_name="Calibri", size_pt=11):
    normal = doc.styles["Normal"]
    normal.font.name      = font_name
    normal.font.size      = Pt(size_pt)
    normal.font.color.rgb = COLOR_DARK_GRAY


def _add_page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    return p


def _add_rule(doc, color=HEX_BLUE, sz="6", space="1",
              space_before=None, space_after=None):
    """Paragraph bottom-border acting as a horizontal rule."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    sz)
    btm.set(qn("w:space"), space)
    btm.set(qn("w:color"), color)
    pBdr.append(btm)
    pPr.append(pBdr)
    p.paragraph_format.space_before = space_before if space_before is not None else Pt(0)
    p.paragraph_format.space_after  = space_after  if space_after  is not None else Pt(6)
    return p


def _insert_field(paragraph, field_code):
    """Inject a PAGE or NUMPAGES field into a paragraph."""
    run = paragraph.add_run()
    fc  = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    run._r.append(fc)

    run2  = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fc2  = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    run3._r.append(fc2)

    for r in (run, run2, run3):
        r.font.name      = "Calibri"
        r.font.size      = Pt(8)
        r.font.color.rgb = COLOR_MID_GRAY


# ── Header / Footer ───────────────────────────────────────────────────────────

def _add_header_footer(doc, report_title, generated_on):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    hdr_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tr = hdr_p.add_run(report_title.upper())
    tr.font.name = "Calibri"; tr.font.size = Pt(8)
    tr.font.bold = True; tr.font.color.rgb = COLOR_PRIMARY

    pPr  = hdr_p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(Pt(468).pt * 20)))
    tabs.append(tab); pPr.append(tabs)

    hdr_p.add_run("\t")
    dr = hdr_p.add_run(generated_on)
    dr.font.name = "Calibri"; dr.font.size = Pt(8); dr.font.color.rgb = COLOR_MID_GRAY

    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single"); btm.set(qn("w:sz"), "4")
    btm.set(qn("w:space"), "1"); btm.set(qn("w:color"), HEX_BLUE)
    pBdr.append(btm); pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    ftr_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pr = ftr_p.add_run("Page ")
    pr.font.name = "Calibri"; pr.font.size = Pt(8); pr.font.color.rgb = COLOR_MID_GRAY
    _insert_field(ftr_p, "PAGE")
    mr = ftr_p.add_run("  of  ")
    mr.font.name = "Calibri"; mr.font.size = Pt(8); mr.font.color.rgb = COLOR_MID_GRAY
    _insert_field(ftr_p, "NUMPAGES")


# ── Cover page ────────────────────────────────────────────────────────────────

def _add_cover_page(doc, report_title, generated_on, total_docs, dept_name=""):
    """Fully prose-based cover page — no tables."""

    # Thick navy top bar
    _add_rule(doc, color=HEX_NAVY, sz="36", space="0",
              space_before=Pt(0), space_after=Pt(0))

    for _ in range(5):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    # Report title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    tr_ = tp.add_run(report_title)
    tr_.font.name = "Calibri"; tr_.font.size = Pt(26)
    tr_.font.bold = True; tr_.font.color.rgb = COLOR_PRIMARY

    # Gold accent rule
    _add_rule(doc, color=HEX_GOLD, sz="12", space="1",
              space_before=Pt(2), space_after=Pt(18))

    # Optional department name
    if dept_name:
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.paragraph_format.space_after = Pt(16)
        dr_ = dp.add_run(dept_name)
        dr_.font.name = "Calibri"; dr_.font.size = Pt(13)
        dr_.font.color.rgb = COLOR_SECONDARY

    # Metadata as centred label: value lines
    for label, value in [
        ("Date Generated",  generated_on),
        ("Total Documents", str(total_docs)),
        ("Classification",  "Official  ·  Internal Use Only"),
    ]:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(2)
        mp.paragraph_format.space_after  = Pt(2)
        lr_ = mp.add_run(f"{label}:  ")
        lr_.font.name = "Calibri"; lr_.font.size = Pt(10)
        lr_.font.bold = True; lr_.font.color.rgb = COLOR_PRIMARY
        vr_ = mp.add_run(value)
        vr_.font.name = "Calibri"; vr_.font.size = Pt(10)
        vr_.font.color.rgb = COLOR_DARK_GRAY

    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    _add_rule(doc, color=HEX_GRAY, sz="4", space="1",
              space_before=Pt(0), space_after=Pt(6))

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr_ = cp.add_run(
        "CONFIDENTIAL  —  This document contains official records. "
        "Unauthorised distribution is strictly prohibited."
    )
    cr_.font.name   = "Calibri"; cr_.font.size = Pt(8)
    cr_.font.italic = True; cr_.font.color.rgb = COLOR_MID_GRAY

    _add_page_break(doc)


# ── Section & entry headings ──────────────────────────────────────────────────

def _section_heading(doc, text):
    """Level-1 task heading with a thick navy rule beneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"; run.font.size = Pt(13)
    run.font.bold = True; run.font.color.rgb = COLOR_PRIMARY
    _add_rule(doc, color=HEX_NAVY, sz="8", space="1",
              space_before=Pt(0), space_after=Pt(8))


def _user_subheading(doc, text):
    """Level-2 employee heading with a fine blue rule beneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"; run.font.size = Pt(11)
    run.font.bold = True; run.font.color.rgb = COLOR_SECONDARY
    _add_rule(doc, color=HEX_BLUE, sz="4", space="1",
              space_before=Pt(0), space_after=Pt(6))


# ── Metadata block (pure prose) ───────────────────────────────────────────────

def _meta_block(doc, doc_data, seq_num=None):
    """Render document metadata as indented label: value prose lines."""
    event_date = doc_data.get("event_date")
    if event_date and hasattr(event_date, "strftime"):
        event_date = event_date.strftime("%B %d, %Y")
    elif event_date:
        event_date = str(event_date)[:10]
    else:
        event_date = "N/A"

    title = doc_data.get("title") or "Untitled Document"
    if seq_num is not None:
        title = f"{seq_num}.  {title}"

    for label, value in [
        ("Document Title", title),
        ("Event Date",     event_date),
        ("Description",    doc_data.get("desc") or "—"),
        ("Submitted By",   doc_data.get("user_name") or "—"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        lr = p.add_run(f"{label}:  ")
        lr.font.name = "Calibri"; lr.font.size = Pt(9.5)
        lr.font.bold = True; lr.font.color.rgb = COLOR_PRIMARY
        vr = p.add_run(str(value))
        vr.font.name = "Calibri"; vr.font.size = Pt(9.5)
        vr.font.color.rgb = COLOR_DARK_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Attachment embed handlers ─────────────────────────────────────────────────

def _embed_image(doc, content_bytes, file_type):
    try:
        img = Image.open(io.BytesIO(content_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG"); buf.seek(0)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        p.add_run().add_picture(buf, width=Inches(5.5))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run("Figure — Attached Image")
        cr.font.name = "Calibri"; cr.font.size = Pt(8)
        cr.font.italic = True; cr.font.color.rgb = COLOR_MID_GRAY

    except Exception as e:
        _error_notice(doc, f"Could not render image: {e}")


def _embed_pdf(doc, content_bytes):
    try:
        pages = convert_from_bytes(content_bytes, dpi=150)
        for i, page in enumerate(pages):
            buf = io.BytesIO()
            page.save(buf, format="PNG"); buf.seek(0)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.add_run().add_picture(buf, width=Inches(5.5))

            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            cr = cap.add_run(f"PDF Attachment  —  Page {i + 1} of {len(pages)}")
            cr.font.name = "Calibri"; cr.font.size = Pt(8)
            cr.font.italic = True; cr.font.color.rgb = COLOR_MID_GRAY

    except Exception as e:
        _error_notice(doc, f"Could not render PDF pages: {e}")


def _embed_docx(doc, content_bytes):
    try:
        from docx import Document as DocxDoc
        inner = DocxDoc(io.BytesIO(content_bytes))

        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after  = Pt(3)
        lr = lp.add_run("Extracted Document Text")
        lr.font.name = "Calibri"; lr.font.size = Pt(9)
        lr.font.bold = True; lr.font.color.rgb = COLOR_SECONDARY

        for para in inner.paragraphs:
            if not para.text.strip():
                continue
            qp = doc.add_paragraph()
            qp.paragraph_format.left_indent  = Inches(0.4)
            qp.paragraph_format.space_before = Pt(1)
            qp.paragraph_format.space_after  = Pt(1)
            qr = qp.add_run(para.text)
            qr.font.name   = "Calibri"; qr.font.size = Pt(9)
            qr.font.italic = True; qr.font.color.rgb = COLOR_DARK_GRAY

            # Left-border quote style via XML
            pPr  = qp._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"),   "single")
            left.set(qn("w:sz"),    "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), HEX_BLUE)
            pBdr.append(left); pPr.append(pBdr)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    except Exception as e:
        _error_notice(doc, f"Could not extract Word document text: {e}")


def _embed_unsupported(doc, file_name, file_type):
    """Prose notice for non-previewable attachment types."""
    for label, value in [
        ("Attachment", file_name or "Unnamed File"),
        ("File Type",  file_type or "unknown"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        lr = p.add_run(f"{label}:  ")
        lr.font.name = "Calibri"; lr.font.size = Pt(10)
        lr.font.bold = True; lr.font.color.rgb = COLOR_PRIMARY
        vr = p.add_run(str(value))
        vr.font.name = "Calibri"; vr.font.size = Pt(10)
        vr.font.color.rgb = COLOR_DARK_GRAY

    note = doc.add_paragraph()
    note.paragraph_format.left_indent  = Inches(0.25)
    note.paragraph_format.space_after  = Pt(10)
    nr = note.add_run(
        "This file type cannot be previewed inline. "
        "Please download the original from the system portal."
    )
    nr.font.name   = "Calibri"; nr.font.size = Pt(9)
    nr.font.italic = True; nr.font.color.rgb = COLOR_MID_GRAY


def _error_notice(doc, message):
    ep = doc.add_paragraph()
    ep.paragraph_format.left_indent = Inches(0.25)
    ep.paragraph_format.space_after = Pt(6)
    er = ep.add_run(f"⚠  {message}")
    er.font.name = "Calibri"; er.font.size = Pt(9)
    er.font.color.rgb = COLOR_RED


# ── Template path ────────────────────────────────────────────────────────────
# The template carries the official header/footer — do not call _add_header_footer.

TEMPLATE_PATH = "excels/template(1).docx"


# ── Main compiler ─────────────────────────────────────────────────────────────

def into_document(documents, report_title="Supporting Documents Report", dept_name=""):
    """
    Build and return a Flask send_file response containing the compiled report.
    Loads excels/template(1).docx so its header/footer are preserved.
    Returns None if the documents list is empty.
    All formatting is pure narrative — no tables used anywhere in the document.
    """
    if not documents:
        return None

    generated_on = datetime.now().strftime("%B %d, %Y")

    # Load the template — inherits header, footer, styles, and margins
    doc = Document(TEMPLATE_PATH)

    # Clear any placeholder content the template ships with, keeping one
    # empty paragraph so the XML stays valid
    body = doc.element.body
    for p in list(body)[:-1]:
        tag = p.tag.split("}")[-1] if "}" in p.tag else p.tag
        if tag in ("p", "tbl"):
            body.remove(p)
    if doc.paragraphs:
        doc.paragraphs[0].clear()

    _set_default_font(doc)

    # Sort: task → user → event_date
    documents.sort(key=lambda x: (
        x.get("task_name") or "Unassigned",
        x.get("user_name") or "",
        str(x.get("event_date") or ""),
    ))

    current_task = None
    current_user = None
    seq_num      = 0

    for doc_data in documents:
        task_name    = doc_data.get("task_name") or "Unassigned Task"
        user_name    = doc_data.get("user_name") or "Unknown User"
        file_type    = (doc_data.get("file_type") or "").lower().strip()
        file_name    = doc_data.get("file_name") or ""
        download_url = doc_data.get("download_url") or ""

        if task_name != current_task:
            if current_task is not None:
                _add_page_break(doc)
            _section_heading(doc, task_name)
            current_task = task_name
            current_user = None

        if user_name != current_user:
            _user_subheading(doc, user_name)
            current_user = user_name

        seq_num += 1
        _meta_block(doc, doc_data, seq_num=seq_num)

        # Fetch attachment from remote storage
        content_bytes = None
        if download_url:
            try:
                resp = requests.get(download_url, timeout=30)
                resp.raise_for_status()
                content_bytes = resp.content
            except Exception as e:
                _error_notice(doc, f"Could not retrieve file: {e}")

        # Render attachment by type
        if content_bytes:
            if file_type in IMAGE_TYPES:
                _embed_image(doc, content_bytes, file_type)
            elif file_type == PDF_TYPE:
                _embed_pdf(doc, content_bytes)
            elif file_type in WORD_TYPES:
                _embed_docx(doc, content_bytes)
            else:
                _embed_unsupported(doc, file_name, file_type)

        # Light grey rule between entries
        _add_rule(doc, color=HEX_GRAY, sz="4", space="4",
                  space_before=Pt(8), space_after=Pt(8))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="Supporting_Documents_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )